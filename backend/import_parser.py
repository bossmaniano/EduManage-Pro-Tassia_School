"""
Import Parser Module for Student Data
Supports .xlsx (Excel) and .docx (Word) file formats
"""

import re
from io import BytesIO

def parse_excel(file_content):
    """
    Parse Excel (.xlsx) files and extract student names
    """
    try:
        import openpyxl
        wb = openpyxl.load_workbook(BytesIO(file_content))
        ws = wb.active
        
        students = []
        headers = []
        
        # First row is headers
        first_row = True
        for row in ws.iter_rows(values_only=True):
            if first_row:
                headers = [str(cell).strip() if cell else '' for cell in row]
                first_row = False
                continue
            
            # Skip empty rows
            if not any(row):
                continue
            
            # Try to find name column - look for common column names
            name = None
            for i, cell in enumerate(row):
                cell_str = str(cell).strip() if cell else ''
                if cell_str and i < len(headers):
                    header_lower = headers[i].lower()
                    if 'name' in header_lower or 'student' in header_lower:
                        name = cell_str
                        break
            
            # If no name column found, use first non-empty cell
            if not name:
                for cell in row:
                    if cell:
                        name = str(cell).strip()
                        break
            
            if name:
                students.append(name)
        
        return students, None
    except Exception as e:
        return [], f"Error parsing Excel file: {str(e)}"


def parse_word(file_content):
    """
    Parse Word (.docx) files and extract student names
    Handles tables and list items
    """
    try:
        from docx import Document
        
        doc = Document(BytesIO(file_content))
        students = []
        
        # First try to extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and len(text) > 1:
                        # Clean up the text - take first line if multiline
                        lines = text.split('\n')
                        name = lines[0].strip()
                        if name and not name.lower().startswith('name'):
                            students.append(name)
        
        # Also check for bullet points/list items
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Check if it's a list item
                if para.style.name.startswith('List'):
                    # Remove bullet characters
                    cleaned = re.sub(r'^[\•\-\*\>\d\.\)]+\s*', '', text)
                    if cleaned and len(cleaned) > 1:
                        students.append(cleaned)
        
        # Remove duplicates while preserving order
        seen = set()
        unique_students = []
        for name in students:
            if name.lower() not in seen:
                seen.add(name.lower())
                unique_students.append(name)
        
        return unique_students, None
    except Exception as e:
        return [], f"Error parsing Word file: {str(e)}"


def sanitize_name(name):
    """
    Convert name to Proper Case and clean up
    """
    if not name:
        return ""
    
    # Convert to title case
    name = name.title()
    
    # Remove extra whitespace
    name = ' '.join(name.split())
    
    # Remove common prefixes/suffixes
    name = re.sub(r'^(Mr\.|Mrs\.|Ms\.|Miss\.|Dr\.|Prof\.)\s*', '', name, flags=re.IGNORECASE)
    
    return name.strip()


def validate_and_process(students_data, existing_students):
    """
    Validate student names and check for duplicates
    Returns: (valid_students, duplicates, invalid_count)
    """
    valid_students = []
    duplicates = []
    invalid_count = 0
    
    # Create set of existing student names (lowercase for comparison)
    existing_names = {s.lower() for s in existing_students}
    
    for name in students_data:
        # Sanitize the name
        cleaned_name = sanitize_name(name)
        
        # Skip empty or very short names
        if not cleaned_name or len(cleaned_name) < 2:
            invalid_count += 1
            continue
        
        # Check for duplicates in import
        if cleaned_name.lower() in [s.lower() for s in valid_students]:
            continue
        
        # Check against existing database
        if cleaned_name.lower() in existing_names:
            duplicates.append(cleaned_name)
            continue
        
        valid_students.append(cleaned_name)
    
    return valid_students, duplicates, invalid_count


def parse_import_file(file_content, filename):
    """
    Main function to parse any supported file format
    """
    ext = filename.lower().split('.')[-1]
    
    if ext == 'xlsx':
        return parse_excel(file_content)
    elif ext == 'docx':
        return parse_word(file_content)
    else:
        return [], f"Unsupported file format: {ext}. Please use .xlsx or .docx files."
